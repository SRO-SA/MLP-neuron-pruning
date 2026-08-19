from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

import math
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REFERENCE = Path(r"C:\Users\srosa\Research\MLP Neuron Pruning\bound_guided_swiglu_moe_pruning_research_draft_v2.docx")
WORKSPACE = Path(r"C:\Users\srosa\Research\MLP Neuron Pruning\MLP Neuron Pruning\qwen_swiglu_pruning")
OUTPUT = WORKSPACE / "bound_guided_swiglu_moe_pruning_research_draft_v3.docx"
ASSET_DIR = WORKSPACE / ".docx_review" / "v3_assets"
EXPECTED_SHA256 = "C853B6FC1FD6ECC4BCA201586282B773C5716F1E15E142321CB878A4E9D09D76"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest().upper()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_exact_table_geometry(table, widths) -> None:
    total_dxa = 10080  # 7.0-inch text width under the retained 0.75-inch margins.
    raw = [float(w) for w in widths]
    dxa = [round(total_dxa * w / sum(raw)) for w in raw]
    dxa[-1] += total_dxa - sum(dxa)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(dxa[i]))
            tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths=None, font_size=8.5, exact_geometry=False) -> None:
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
        if exact_geometry:
            apply_exact_table_geometry(table, widths)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if ri == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
                    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
                    run.font.size = Pt(font_size)
                    if ri == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        if ri == 0:
            set_repeat_table_header(row)
            for cell in row.cells:
                set_cell_shading(cell, "4F81BD")


def add_rows_to_existing(table, rows) -> None:
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
    # Preserve the source table's equal-width grid and visual language.
    widths = [7.0 / len(table.columns)] * len(table.columns)
    format_table(table, widths=widths, font_size=8.5)


def find_paragraph(doc, exact=None, starts=None):
    for p in doc.paragraphs:
        if exact is not None and p.text == exact:
            return p
        if starts is not None and p.text.startswith(starts):
            return p
    raise ValueError(f"Paragraph not found: exact={exact!r} starts={starts!r}")


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def equation_after(paragraph, text: str):
    p = paragraph_after(paragraph, style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Cambria Math")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Cambria Math")
    run.font.size = Pt(10.5)
    return p


def note_after(paragraph, label: str, text: str):
    p = paragraph_after(paragraph, style="Normal")
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF2F8")
    p_pr.append(shd)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(23, 54, 93)
    p.add_run(text)
    return p


class BeforeBuilder:
    def __init__(self, doc, target_paragraph):
        self.doc = doc
        self.target = target_paragraph

    def p(self, text="", style="Normal"):
        p = self.target.insert_paragraph_before(text, style=style)
        return p

    def equation(self, text):
        p = self.p("", "Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Cambria Math"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Cambria Math")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Cambria Math")
        return p

    def table(self, headers, rows, widths, font_size=8.2):
        table = self.doc.add_table(rows=1, cols=len(headers))
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)
        for values in rows:
            cells = table.add_row().cells
            for i, value in enumerate(values):
                cells[i].text = str(value)
        format_table(table, widths=widths, font_size=font_size, exact_geometry=True)
        self.target._p.addprevious(table._tbl)
        return table

    def picture(self, image_path: Path, width=6.2):
        p = self.p("", "Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(width))
        return p


def make_figures() -> tuple[Path, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    selector_path = ASSET_DIR / "target6_ranking_comparison.png"
    tightness_path = ASSET_DIR / "bound_tightness.png"

    font_path = r"C:\Windows\Fonts\times.ttf"
    bold_path = r"C:\Windows\Fonts\timesbd.ttf"
    font = ImageFont.truetype(font_path, 28)
    small = ImageFont.truetype(font_path, 24)
    title_font = ImageFont.truetype(bold_path, 34)
    label_font = ImageFont.truetype(bold_path, 25)

    def centered(draw, xy, text, font_obj, fill="#222222"):
        box = draw.textbbox((0, 0), text, font=font_obj)
        draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1]), text, font=font_obj, fill=fill)

    # Target-6 grouped bar chart.
    img = Image.new("RGB", (1500, 880), "white")
    draw = ImageDraw.Draw(img)
    centered(draw, (750, 28), "Target 6%: Fixed RMSNorm Allocation", title_font, "#17365D")
    centered(draw, (750, 72), "Ranking-only comparison", font, "#444444")
    left, top, right, bottom = 165, 140, 1430, 720
    draw.line((left, top, left, bottom), fill="#333333", width=3)
    draw.line((left, bottom, right, bottom), fill="#333333", width=3)
    ymax = 3.2
    for tick in [0, 0.5, 1.0, 1.5, 2.0, 2.5, 3.0]:
        y = bottom - (tick / ymax) * (bottom - top)
        draw.line((left, y, right, y), fill="#D9E2F3", width=2)
        draw.text((80, y - 14), f"{tick:.1f}", font=small, fill="#333333")
    labels = ["Original bound", "Activation", "Ellipsoid"]
    wiki = [2.979, 1.801, 0.610]
    c4 = [1.901, 2.510, 1.672]
    group_w = (right - left) / 3
    bar_w = 115
    for i, label in enumerate(labels):
        cx = left + group_w * (i + 0.5)
        for value, dx, color in [(wiki[i], -bar_w - 8, "#4F81BD"), (c4[i], 8, "#C0504D")]:
            x0 = cx + dx
            x1 = x0 + bar_w
            y0 = bottom - (value / ymax) * (bottom - top)
            draw.rectangle((x0, y0, x1, bottom), fill=color)
            centered(draw, ((x0 + x1) / 2, y0 - 33), f"{value:.3f}%", small)
        centered(draw, (cx, bottom + 22), label, small)
    draw.multiline_text((12, 365), "Relative PPL\nincrease (%)", font=small, fill="#333333", spacing=5, align="center")
    draw.rectangle((990, 112, 1025, 137), fill="#4F81BD")
    draw.text((1035, 107), "WikiText-2", font=small, fill="#333333")
    draw.rectangle((1220, 112, 1255, 137), fill="#C0504D")
    draw.text((1265, 107), "C4", font=small, fill="#333333")
    img.save(selector_path, dpi=(180, 180))

    # Log-scale tightness chart.
    img = Image.new("RGB", (1500, 880), "white")
    draw = ImageDraw.Draw(img)
    centered(draw, (750, 28), "Empirical Certificate Tightness", title_font, "#17365D")
    centered(draw, (750, 72), "Observed contribution divided by bound (log scale)", font, "#444444")
    left, top, right, bottom = 185, 140, 1430, 720
    draw.line((left, top, left, bottom), fill="#333333", width=3)
    draw.line((left, bottom, right, bottom), fill="#333333", width=3)
    log_min, log_max = -5.0, 0.0
    for exponent in [-5, -4, -3, -2, -1, 0]:
        y = bottom - ((exponent - log_min) / (log_max - log_min)) * (bottom - top)
        draw.line((left, y, right, y), fill="#D9E2F3", width=2)
        draw.text((82, y - 14), f"10^{exponent}", font=small, fill="#333333")
    metrics = ["Median", "p95", "p99", "Maximum"]
    sphere = [2.9067920877423603e-05, 1.78006975329481e-04, 4.943700262811039e-04, 3.276412561535835e-02]
    ellip = [1.7678930453257635e-04, 8.960256003774701e-04, 2.009214940480889e-03, 1.7312145233154297e-01]
    group_w = (right - left) / 4
    bar_w = 90
    for i, label in enumerate(metrics):
        cx = left + group_w * (i + 0.5)
        for value, dx, color in [(sphere[i], -bar_w - 7, "#9EADBA"), (ellip[i], 7, "#4F81BD")]:
            val_log = max(log_min, math.log10(value))
            y0 = bottom - ((val_log - log_min) / (log_max - log_min)) * (bottom - top)
            x0 = cx + dx
            x1 = x0 + bar_w
            draw.rectangle((x0, y0, x1, bottom), fill=color)
        centered(draw, (cx, bottom + 22), label, small)
    draw.rectangle((980, 112, 1015, 137), fill="#9EADBA")
    draw.text((1025, 107), "Sphere", font=small, fill="#333333")
    draw.rectangle((1180, 112, 1215, 137), fill="#4F81BD")
    draw.text((1225, 107), "Ellipsoid", font=small, fill="#333333")
    centered(draw, (750, 800), "Higher values below 1 indicate a tighter valid bound", label_font, "#444444")
    img.save(tightness_path, dpi=(180, 180))
    return selector_path, tightness_path


def main() -> None:
    if sha256(REFERENCE) != EXPECTED_SHA256:
        raise RuntimeError("Reference DOCX hash changed; fresh template distillation is required.")

    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    original_tables = list(doc.tables)

    selector_figure, tightness_figure = make_figures()

    # Title block and current summary.
    replace_paragraph_text(doc.paragraphs[0], "Certified Ellipsoidal Bound-Guided Structured SwiGLU and MoE Expert-Channel Pruning\nwith Residual Reconstruction for Efficient LLM Inference")
    replace_paragraph_text(doc.paragraphs[1], "Working Conference-Style Research Draft — Version 3 Extended Internal Version")
    replace_paragraph_text(doc.paragraphs[3], "August 2026")
    replace_paragraph_text(
        find_paragraph(doc, starts="Large language model inference is strongly affected"),
        "Large language model inference is strongly affected by transformer feed-forward blocks and, in mixture-of-experts (MoE) models, by the expert MLPs activated for each token. This Version 3 draft preserves the complete dense and MoE pruning study from earlier versions and expands its mathematical and empirical core. The original method physically removes coupled SwiGLU gate/up/down channels using a static RMSNorm-inspired spherical score and optionally repairs the remaining down projection through residual reconstruction. The new expansion derives an exact ellipsoidal upper bound induced by the coordinatewise RMSNorm scale vector, separates cross-layer budget allocation from within-layer channel ranking, and evaluates these choices independently. On Qwen3-30B-A3B, controlled target-4% experiments show that ellipsoid ranking improves both RMSNorm-bound and down-norm allocation plans on WikiText-2 and C4. At an actual 6.207% expert-width reduction under a fixed RMSNorm allocation, ellipsoid ranking yields relative perplexity increases of 0.610% on WikiText-2 and 1.672% on C4, compared with 2.979%/1.901% for the original bound ranking and 1.801%/2.510% for activation ranking. Paired bootstrap confidence intervals favor ellipsoid ranking on both datasets. A sampled certificate audit finds zero numerical violations and shows that the ellipsoid bound is approximately four to six times tighter than the spherical bound by observed-contribution ratios, although both remain conservative. Earlier dense Qwen2.5 and MoE residual-reconstruction results are retained in full. The evidence supports a calibration-free, certificate-derived within-layer selector, while submission-quality claims still require downstream tasks, matched external baselines, maximum-across-expert aggregation, and measured end-to-end serving gains."
    )
    replace_paragraph_text(
        find_paragraph(doc, starts="Keywords:"),
        "Keywords: LLM compression; structured pruning; certified pruning; RMSNorm ellipsoid; SwiGLU; mixture of experts; expert-channel pruning; residual reconstruction; calibration-free ranking; Qwen; inference efficiency."
    )

    for section in doc.sections:
        for p in section.footer.paragraphs:
            if p.text:
                replace_paragraph_text(p, "Certified Ellipsoidal SwiGLU and MoE Expert-Channel Pruning — Version 3 Working Draft")

    # Preserve and extend project status and related work tables.
    add_rows_to_existing(original_tables[0], [
        ("Exact RMSNorm ellipsoid expansion", "Derived and implemented", "Expert-channel bound uses the full coordinatewise RMSNorm scale rather than a gamma-infinity sphere."),
        ("Allocation/ranking separation", "Implemented and evaluated", "Fixed-plan experiments isolate cross-layer budget allocation from within-layer channel choice."),
        ("Paired statistical evaluation", "Implemented at target 6%", "10,000-resample paired document bootstrap comparisons on WikiText-2 and C4."),
        ("Certificate tightness audit", "Implemented", "3,283,968 sampled expert-channel evaluations; zero reported numerical violations."),
    ])
    add_rows_to_existing(original_tables[1], [
        ("HEAPr [13]", "Hessian-based atomic-expert pruning using output-space second-order information.", "Directly relevant fine-grained MoE baseline; our selector is weight-only and calibration-free but currently targets much smaller pruning ratios."),
        ("CAMERA [14]", "Training-free micro-expert redundancy analysis with structured pruning and quantization.", "Very close pruning unit; Version 3 novelty must rest on the explicit RMSNorm ellipsoid certificate and allocation/ranking analysis."),
        ("Attribution-guided structural MoE compression [15]", "Channel attribution and coverage-maximized allocation for heterogeneous structural pruning.", "Stronger allocation and compression baseline; our present advantage is the formal local bound, not demonstrated high-ratio state of the art."),
    ])

    # Extend contribution list without deleting earlier contributions.
    anchor = find_paragraph(doc, exact="A conservative research positioning relative to SparseGPT, Wanda, LLM-Pruner, SliceGPT, SlimLLM, and related structured pruning work.")
    for text in [
        "An exact RMSNorm-induced ellipsoid bound for individual SwiGLU channel contributions, together with a short spectral proof.",
        "A conceptual and experimental separation between cross-layer pruning-budget allocation and within-layer channel ranking.",
        "Controlled fixed-allocation comparisons showing that ellipsoid ranking improves the original bound and activation ranking at the 6% operating point on both WikiText-2 and C4.",
        "An empirical certificate audit reporting zero sampled violations and substantially tighter observed-contribution ratios than the spherical bound.",
    ]:
        anchor = paragraph_after(anchor, text, "List Bullet")

    # Updated positioning after the existing conservative novelty statement.
    anchor = find_paragraph(doc, starts="Conservative novelty statement.")
    anchor = paragraph_after(
        anchor,
        "Version 3 positioning. Fine-grained MoE pruning is now clearly occupied by atomic-expert, micro-expert, and channel-attribution methods such as HEAPr, CAMERA, and attribution-guided structural compression [13–15]. Consequently, Version 3 does not claim novelty for expert-channel pruning itself. Its narrower claim is a calibration-free, architecture-derived upper bound that follows exactly from the RMSNorm feasible set, plus evidence that this bound improves within-layer channel ranking when allocation is held fixed. The current pruning ratios are substantially below the aggressive ratios reported by several recent competitors, so competitiveness must be established through matched experiments rather than cross-paper headline comparisons."
    )

    # Mathematical expansion after the original score section.
    anchor = find_paragraph(doc, starts="The implemented angle-aware version replaces")
    anchor = paragraph_after(anchor, "3.2.1 Intuition: From an RMSNorm Sphere to Its Actual Ellipsoid", "Heading 3")
    anchor = paragraph_after(
        anchor,
        "The original argument begins with a safe Euclidean radius. Write the RMSNorm output as r = Γq, where Γ = diag(γ) is the learned coordinatewise scale and q = x / sqrt(||x||²/d + ε). Then ||q|| < sqrt(d). Replacing Γ by its largest diagonal magnitude gives ||r|| ≤ sqrt(d)||γ||∞. This encloses every feasible normalized input in a sphere. It is mathematically safe, and the resulting factor is constant within a layer, but it imagines that all coordinates can simultaneously receive the largest RMSNorm amplification."
    )
    anchor = equation_after(anchor, "Bᵢˢᵖʰ = [d ||γ||∞² / 2] (||gᵢ|| ||uᵢ|| + |gᵢᵀuᵢ|) ||dᵢ||")
    anchor = paragraph_after(
        anchor,
        "Geometrically, RMSNorm does not produce a sphere after the learned γ scaling. It produces the image of a radius-sqrt(d) ball under Γ: an axis-aligned ellipsoid. Coordinates with small |γ_j| cannot reach the radius permitted by ||γ||∞, whereas coordinates with large |γ_j| can. The ellipsoid therefore contains the same feasible inputs without paying the largest amplification in every direction."
    )
    anchor = paragraph_after(anchor, "3.2.2 Exact Ellipsoid Bound", "Heading 3")
    anchor = paragraph_after(
        anchor,
        "For channel i, define the RMSNorm-scaled gate and up vectors aᵢ = Γgᵢ and bᵢ = Γuᵢ. Because |SiLU(t)| ≤ |t|, the norm of the deleted channel contribution is bounded by |qᵀaᵢ| |qᵀbᵢ| ||dᵢ||. The remaining maximization is exactly solvable over ||q|| ≤ sqrt(d):"
    )
    anchor = equation_after(anchor, "max ||q||≤√d |qᵀaᵢ| |qᵀbᵢ| = d/2 (||aᵢ|| ||bᵢ|| + |aᵢᵀbᵢ|)")
    anchor = equation_after(anchor, "Bᵢᵉˡˡ = d/2 (||Γgᵢ|| ||Γuᵢ|| + |(Γgᵢ)ᵀ(Γuᵢ)|) ||dᵢ||")
    anchor = paragraph_after(
        anchor,
        "The angle term is not an ad hoc correction. The product (qᵀa)(qᵀb) is the quadratic form qᵀ(abᵀ + baᵀ)q/2. The two nonzero eigenvalues of this symmetric matrix are (aᵀb ± ||a||||b||)/2. Maximizing the absolute Rayleigh quotient over a radius-sqrt(d) ball gives the expression above. The ellipsoid bound is therefore exact for the bilinear relaxation produced by |SiLU(t)| ≤ |t|. It is still conservative because the SiLU envelope and independent channel-wise triangle bounds do not exploit the actual routed-token distribution or cancellation between removed channels."
    )
    anchor = note_after(
        anchor,
        "Mathematical intuition.",
        "The spherical score asks what would happen if every direction enjoyed the largest RMSNorm scale. The ellipsoid score first stretches each gate/up vector by its actual coordinatewise γ and then asks for the worst direction. A channel that relies on strongly suppressed coordinates becomes provably less capable of producing a large output."
    )
    anchor = paragraph_after(anchor, "3.2.3 Local Ranking Is Not Global Layer Allocation", "Heading 3")
    anchor = paragraph_after(
        anchor,
        "A local channel bound answers: within this expert and layer, how large could the direct deleted contribution be? It does not answer how a perturbation at one layer is amplified by all downstream layers. Raw ellipsoid magnitudes also inherit layer-specific γ and weight scales. The exploratory 2×2 experiment showed that ellipsoid ranking improved channel choice inside a fixed layer budget, while ellipsoid-based global allocation concentrated pruning in harmful layers. Version 3 therefore separates two variables: an allocation source chooses the aligned count k_l removed from each layer, and a ranking source chooses the k_l channel identities within that layer. This separation is a methodological correction, not a weakening of the theorem."
    )
    anchor = paragraph_after(anchor, "3.2.4 MoE Aggregation and Scope of the Certificate", "Heading 3")
    anchor = paragraph_after(
        anchor,
        "For expert e, layer l, and channel i, B_{l,e,i}^{ell} is an expert-specific bound. The current packed same-channel implementation aggregates expert scores using the 95th percentile before selecting a shared channel ID. This p95 statistic is a robust heuristic and is not a worst-case certificate across all experts. Replacing p95 by max_e B_{l,e,i}^{ell} would provide the conservative shared-channel score needed for an all-expert statement. For a set P of removed channels in one expert, the triangle inequality gives ||ΔE(r)|| ≤ Σ_{i∈P} B_i^{ell}. If router weights are nonnegative and normalized, the routed mixture error is at most their weighted sum of expert error bounds. These are valid but potentially loose group guarantees."
    )
    anchor = paragraph_after(anchor, "3.2.5 Extension Beyond SwiGLU: Capped GLUs and SiTU-GLU", "Heading 3")
    anchor = paragraph_after(
        anchor,
        "The ellipsoid geometry is not specific to SiLU. It applies whenever scalar envelopes are available for the two multiplicative branches. Kimi K3, for example, uses SiTU-GLU with gate branch β₁ tanh(x/β₁)σ(x) and up branch β₂ tanh(x/β₂), with β₁ = 4 and β₂ = 25 [19]. Since each branch is both linearly bounded near the origin and globally capped, a channel admits the valid composite bound"
    )
    anchor = equation_after(anchor, "Bᵢˢⁱᵀᵁ = ||dᵢ|| min{β₁β₂, β₁√d||Γuᵢ||, β₂√d||Γgᵢ||, d/2(||Γgᵢ||||Γuᵢ|| + |(Γgᵢ)ᵀ(Γuᵢ)|)}")
    anchor = paragraph_after(
        anchor,
        "This bound can be tighter than the SwiGLU form because saturation supplies the input-independent cap β₁β₂ = 100. Kimi K3 already uses partial expert activation—16 of 896 routed experts per token—to reduce active computation, but routing sparsity does not reduce the stored width of every expert. Structured channel pruning remains complementary: it reduces parameters and per-active-expert matrix dimensions. This extension is theoretical in Version 3 and has not been evaluated on the full Kimi K3 checkpoint."
    )

    # Reproducibility additions.
    anchor = find_paragraph(doc, starts="The selected full benchmark script was derived")
    anchor = paragraph_after(
        anchor,
        "Version 3 adds an explicit allocation/ranking experiment interface. Every hybrid run is validated against the allocation source's per-layer count vector, starts from a freshly loaded unpruned checkpoint, preserves the requested alignment and layer caps, and records allocation source, ranking source, actual removed layer-channels, removed expert-neurons, token counts, and result paths. This prevents the ranking comparison from being confounded by cumulative pruning or different layer budgets."
    )
    anchor = paragraph_after(
        anchor,
        "For the target-6 primary experiment, per-document negative log-likelihood values were retained and paired bootstrap comparisons used 10,000 resamples over the same 1,024 evaluation documents. This is preferable to treating deterministic reruns as independent samples. The bound audit separately sampled 4,276 experts, 25,376 routed inputs, and 3,283,968 expert-channel contributions."
    )

    # Version 3 experimental protocol before dense historical results.
    target = find_paragraph(doc, exact="6. Dense Qwen2.5 Results")
    b = BeforeBuilder(doc, target)
    b.p("5.1 Version 3 Controlled Allocation/Ranking Protocol", "Heading 2")
    b.p(
        "The new Qwen3-30B-A3B experiments use two independent selector roles. The allocation selector determines an aligned per-layer removal-count vector; the ranking selector chooses channel identities within each layer. The principal comparison fixes RMSNorm-bound allocation and substitutes original-bound, activation, or ellipsoid ranking. A secondary comparison fixes down-norm allocation and substitutes down-norm or ellipsoid ranking. Unless stated otherwise, expert scores are aggregated by p95 across the 128 experts sharing a channel ID."
    )
    b.p(
        "The target-4 and target-6 experiments evaluate 1,024 documents. WikiText-2 contains 170,564 evaluated tokens. C4 contains 285,750 evaluated tokens. Baseline PPL values differ slightly between target groups because the current artifacts were produced by separate runs; comparisons within a row group use the same baseline and exact samples. These Version 3 results must not be numerically merged with the n_eval=512 Version 2 benchmark, whose baselines were 12.3194 on WikiText-2 and 14.4115 on C4."
    )
    b.p(
        "Activation ranking is calibration-dependent; the bound rankings are weight-only after model loading. The Version 3 interpretation therefore distinguishes predictive performance from calibration cost. Final reporting must also confirm calibration/evaluation disjointness for every activation-based run."
    )

    # Historical labels, preserving the older result narratives.
    anchor = find_paragraph(doc, starts="The dense experiments validate the original idea")
    paragraph_after(anchor, "Version 3 retention note. These dense tables are preserved from the original implementation and have not yet been rerun with ellipsoid ranking. They remain evidence for physical target pruning and residual reconstruction, not a direct evaluation of the new selector.", "Normal")
    anchor = find_paragraph(doc, starts="The main new result is the selected full MoE benchmark")
    paragraph_after(anchor, "Version 3 retention note. Section 7.1 reports the complete Version 2 n_eval=512 MoE residual benchmark. Section 7.2 adds the later n_eval=1024 allocation/ranking experiments. Because baseline PPL values changed, the two protocols are interpreted separately.", "Normal")

    # New results inserted before existing Section 8.
    target = find_paragraph(doc, exact="8. What We Have Achieved So Far")
    b = BeforeBuilder(doc, target)
    b.p("7.2 Version 3 Ellipsoid Allocation/Ranking Results", "Heading 2")
    b.p(
        "The first diagnostic crossed the original and ellipsoid choices for global allocation and within-layer ranking. The result localized the failure of the first ellipsoid implementation: ellipsoid ranking was beneficial under the original allocation, whereas ellipsoid allocation was harmful under either ranking. The negative WikiText-2 change in one exploratory row is treated as no measurable degradation rather than a general quality improvement."
    )
    b.p("Table V3-1. Exploratory 2% allocation/ranking diagnostic on WikiText-2.", "Caption")
    b.table(
        ["Allocation", "Ranking", "Actual", "Layer ch.", "Base PPL", "Pruned PPL", "Rel. change"],
        [
            ("RMSNorm bound", "RMSNorm bound", "≈2.3%", "832", "12.3194", "12.3699", "+0.410%"),
            ("RMSNorm bound", "Ellipsoid", "2.257%", "832", "12.3194", "12.2125", "−0.868%"),
            ("Ellipsoid", "Ellipsoid", "2.083%", "768", "12.3194", "12.5426", "+1.812%"),
            ("Ellipsoid", "RMSNorm bound", "2.083%", "768", "12.3194", "12.6314", "+2.532%"),
        ],
        [1.15, 1.15, 0.75, 0.72, 0.95, 0.98, 1.05],
    )
    b.p(
        "At target 4%, ellipsoid ranking again improves both tested allocation plans. Under RMSNorm allocation it changes WikiText-2 degradation from +2.174% to −0.414% and C4 from +1.345% to +1.087%. Under down-norm allocation it changes WikiText-2 from +1.120% to −0.418% and C4 from +1.637% to +1.236%. The two allocation plans remove different totals, so their absolute rows do not establish which allocation policy is superior."
    )
    b.p("Table V3-2. Target-4 controlled allocation/ranking results (n_eval=1024).", "Caption")
    b.table(
        ["Experiment", "Data", "Alloc.", "Rank", "Actual", "Ch.", "Base", "Pruned", "Rel."],
        [
            ("rms/rms", "WT2", "RMS", "RMS", "4.340%", "1600", "12.8340", "13.1130", "+2.174%"),
            ("rms/ellip", "WT2", "RMS", "Ellip", "4.340%", "1600", "12.8340", "12.7809", "−0.414%"),
            ("down/down", "WT2", "Down", "Down", "4.167%", "1536", "12.8340", "12.9777", "+1.120%"),
            ("down/ellip", "WT2", "Down", "Ellip", "4.167%", "1536", "12.8340", "12.7803", "−0.418%"),
            ("rms/rms", "C4", "RMS", "RMS", "4.340%", "1600", "14.5607", "14.7565", "+1.345%"),
            ("rms/ellip", "C4", "RMS", "Ellip", "4.340%", "1600", "14.5607", "14.7190", "+1.087%"),
            ("down/down", "C4", "Down", "Down", "4.167%", "1536", "14.5607", "14.7990", "+1.637%"),
            ("down/ellip", "C4", "Down", "Ellip", "4.167%", "1536", "14.5607", "14.7407", "+1.236%"),
        ],
        [1.05, 0.5, 0.55, 0.58, 0.75, 0.55, 0.8, 0.8, 0.72],
        7.7,
    )
    b.p("7.3 Target-6 Fixed-Allocation Attribution", "Heading 2")
    b.p(
        "The target-6 primary experiment is the strongest current method-attribution result. Under the same RMSNorm allocation, identical 2,288 removed layer-channels, identical expert aggregation, and identical evaluation documents, ellipsoid ranking outperforms both the original bound and activation ranking on both corpora. The secondary down-norm allocation removes 2,256 channels and gives a similar C4 result, supporting—but not proving—allocation robustness."
    )
    b.p("Table V3-3. Target-6 controlled allocation/ranking results (n_eval=1024).", "Caption")
    b.table(
        ["Experiment", "Data", "Alloc.", "Rank", "Actual", "Ch.", "Base", "Pruned", "Rel."],
        [
            ("rms/rms", "WT2", "RMS", "RMS", "6.207%", "2288", "12.8342", "13.2165", "+2.979%"),
            ("rms/act", "WT2", "RMS", "Act", "6.207%", "2288", "12.8342", "13.0653", "+1.801%"),
            ("rms/ellip", "WT2", "RMS", "Ellip", "6.207%", "2288", "12.8342", "12.9125", "+0.610%"),
            ("down/ellip", "WT2", "Down", "Ellip", "6.120%", "2256", "12.8342", "12.9549", "+0.941%"),
            ("rms/rms", "C4", "RMS", "RMS", "6.207%", "2288", "14.5643", "14.8411", "+1.901%"),
            ("rms/act", "C4", "RMS", "Act", "6.207%", "2288", "14.5643", "14.9298", "+2.510%"),
            ("rms/ellip", "C4", "RMS", "Ellip", "6.207%", "2288", "14.5643", "14.8078", "+1.672%"),
            ("down/ellip", "C4", "Down", "Ellip", "6.120%", "2256", "14.5643", "14.8068", "+1.665%"),
        ],
        [1.05, 0.5, 0.55, 0.58, 0.75, 0.55, 0.8, 0.8, 0.72],
        7.7,
    )
    b.p("Figure V3-1. Target-6 relative perplexity increase under a fixed RMSNorm allocation.", "Caption")
    b.picture(selector_figure)
    b.p("7.4 Paired Statistical Comparisons", "Heading 2")
    b.p(
        "Per-document NLL differences were paired because all rankings were evaluated on the same documents. Negative ellipsoid-minus-competitor NLL means ellipsoid ranking is better. All four 95% bootstrap intervals exclude zero. The C4 advantage over the original bound is smaller than the other effects but remains negative in the reported interval."
    )
    b.p("Table V3-4. Paired target-6 ellipsoid-versus-competitor NLL comparisons.", "Caption")
    b.table(
        ["Dataset", "Competitor", "Mean NLL Δ", "95% lower", "95% upper", "Docs", "Tokens", "Resamples"],
        [
            ("WikiText-2", "Original bound", "−0.023276", "−0.026276", "−0.020332", "1024", "170,564", "10,000"),
            ("WikiText-2", "Activation", "−0.011767", "−0.015570", "−0.007894", "1024", "170,564", "10,000"),
            ("C4", "Original bound", "−0.002246", "−0.004448", "−0.000238", "1024", "285,750", "10,000"),
            ("C4", "Activation", "−0.008206", "−0.010625", "−0.005849", "1024", "285,750", "10,000"),
        ],
        [1.05, 1.05, 0.9, 0.9, 0.9, 0.65, 0.8, 0.75],
        8.0,
    )
    b.p("7.5 Certificate Validity and Tightness", "Heading 2")
    b.p(
        "The tightness metric is observed channel-contribution norm divided by its bound. A larger ratio below one means a tighter certificate. Across 3,283,968 sampled expert-channel evaluations, neither bound reported a numerical violation. The ellipsoid ratio is approximately 6.08× larger at the median, 5.03× at p95, 4.07× at p99, and 5.28× at the maximum. The maximum observed contribution reaches 17.31% of the ellipsoid bound versus 3.28% of the spherical bound. Nevertheless, the median and p95 ratios remain below 0.1%, demonstrating that a uniform channel-wise certificate is still highly conservative on typical routed inputs."
    )
    b.p("Table V3-5. Sampled expert-channel bound tightness.", "Caption")
    b.table(
        ["Bound", "Count", "Median", "p95", "p99", "Maximum", "Violations"],
        [
            ("Ellipsoid", "3,283,968", "0.0001768", "0.0008960", "0.0020092", "0.1731215", "0"),
            ("Sphere", "3,283,968", "0.0000291", "0.0001780", "0.0004944", "0.0327641", "0"),
        ],
        [1.0, 1.0, 1.0, 1.0, 1.0, 1.05, 0.95],
    )
    b.p("Figure V3-2. Observed contribution divided by bound; higher values below one indicate tighter certificates.", "Caption")
    b.picture(tightness_figure)
    b.p("7.6 Comparison with the Version 2 MoE Benchmark", "Heading 2")
    b.p(
        "The Version 2 6.2% pure-deletion rows reported +3.8% relative PPL on WikiText-2 and +2.2% on C4 at n_eval=512. The Version 3 ellipsoid rows report +0.610% and +1.672% at n_eval=1024. These values are directionally encouraging but are not a controlled before/after comparison: the baseline PPL changed from 12.3194/14.4115 to 12.8342/14.5643, establishing that the evaluated sample sets or protocol differ. The valid causal comparison is the contemporaneous fixed-allocation result within Version 3: ellipsoid versus original-bound and activation ranking under identical evaluation documents."
    )
    b.p("7.7 Current Empirical Conclusion", "Heading 2")
    b.p(
        "The combined 2%, 4%, and 6% evidence supports three claims. First, coordinatewise RMSNorm geometry contains useful channel-ranking information beyond the original spherical proxy. Second, this advantage persists under two allocation sources and two language-modeling corpora. Third, raw local-bound magnitudes should not be used as a global layer-sensitivity measure without normalization or downstream amplification modeling. The current evidence does not yet establish superiority at high pruning ratios, downstream-task preservation, or realized inference speed."
    )

    # Extend achievements, claims, and next-experiment tables without removing existing rows.
    add_rows_to_existing(original_tables[8], [
        ("Exact ellipsoid theorem", "Derived the exact bilinear maximum over the RMSNorm-induced ellipsoid and implemented the resulting selector."),
        ("Controlled selector attribution", "Separated allocation and ranking; ellipsoid ranking improves fixed-plan baselines at targets 4% and 6%."),
        ("Statistical support", "Paired target-6 NLL intervals favor ellipsoid ranking over original-bound and activation ranking on both datasets."),
        ("Certificate audit", "Zero sampled violations and roughly 4–6× tighter observed/bound ratios than the sphere."),
    ])
    add_rows_to_existing(original_tables[9], [
        ("Ellipsoid ranking is better than the original ranking under fixed target-6 allocation", "Yes, for current WT2/C4 protocol", "Paired 95% NLL intervals exclude zero on both datasets."),
        ("Ellipsoid ranking is calibration-free", "Yes for scoring", "Confirm that no evaluation data enter model selection or allocation tuning."),
        ("Current p95 packed-channel method is uniformly certified across experts", "No", "Run max-expert aggregation or state p95 as a heuristic."),
        ("The ellipsoid selector is state of the art at high pruning ratios", "Not yet", "Need target 8%, downstream tasks, and matched HEAPr/CAMERA/attribution baselines."),
    ])
    add_rows_to_existing(original_tables[10], [
        ("Target-8 fixed-allocation ranking", "Compare original, activation, and ellipsoid rankings under one exact plan on WT2/C4.", "Defines the failure curve of the new selector."),
        ("p95 versus max aggregation", "Evaluate robust heuristic versus conservative all-expert score at 4% and 6%.", "Separates best empirical quality from the strongest certificate statement."),
        ("Dense ellipsoid validation", "Rerun Qwen2.5-3B/7B at 4% and 6% with original, activation, and ellipsoid rankings.", "Supports the paper's general SwiGLU claim."),
        ("Paired downstream evaluation", "Evaluate selected 4%/6% checkpoints on HellaSwag, ARC, MMLU subset, and GSM8K where feasible.", "Moves beyond perplexity."),
        ("Matched modern MoE baselines", "Compare against or carefully reproduce HEAPr, CAMERA, and attribution-guided compression.", "Required for a current novelty and competitiveness claim."),
    ])

    # Add Version 3 positioning paragraphs.
    anchor = find_paragraph(doc, starts="The strongest differentiator is the MoE expert-channel result")
    anchor = paragraph_after(
        anchor,
        "Version 3 changes the strongest differentiator. Physical expert-channel pruning alone is no longer sufficiently distinctive given recent atomic- and micro-expert methods. The more defensible contribution is the exact RMSNorm ellipsoid analysis and the controlled evidence that a certificate-derived, weight-only ranking can outperform a data-dependent activation ranking at the tested 6% operating point. This must still be compared against the closest modern methods at matched pruning units and budgets."
    )
    anchor = paragraph_after(
        anchor,
        "The allocation/ranking diagnostic is itself useful: it shows that a mathematically tighter local bound need not be a better global allocation score. This clarifies a common conceptual error in structured pruning. A future global objective should combine the local certificate with a layer-sensitivity or downstream-amplification factor rather than treating raw bounds from different layers as commensurate."
    )

    # Add new limitations after the last existing limitation bullet.
    anchor = find_paragraph(doc, starts="The C4/WikiText rows in the current summary")
    for text in [
        "The ellipsoid bound is substantially tighter than the sphere but remains very conservative on typical routed inputs; median observed contributions are far below the certified maximum.",
        "The current p95 expert aggregation is a robust ranking heuristic, not a uniform certificate across every expert. Maximum aggregation remains to be evaluated.",
        "The Version 2 and Version 3 evaluations use different sample counts and baseline PPL values. Their headline percentages are directional comparisons only; causal claims rely on within-protocol fixed-plan experiments.",
        "The new selector has not yet been evaluated at 8%, on downstream tasks, or against current fine-grained MoE methods at matched budgets.",
        "Same-channel-across-experts pruning preserves a simple packed layout but is more restrictive than equal-width expert-specific repacking. Expert-specific channel IDs remain an important future extension.",
    ]:
        anchor = paragraph_after(anchor, text, "List Bullet")

    # Add a Version 3 execution paragraph after the next-experiment table caption area.
    anchor = find_paragraph(doc, exact="Table 11. Recommended experiments before submission.")
    paragraph_after(
        anchor,
        "Version 3 priority order. First complete the target-8 fixed-allocation selector comparison and p95-versus-max aggregation. Second evaluate downstream accuracy and actual serving performance for the selected 4% and 6% checkpoints. Third add a matched modern MoE pruning baseline and dense ellipsoid validation. Exact two-dimensional SiLU optimization and layer-sensitivity modeling are promising extensions, but they should follow these empirical gates so that the paper's central claim remains falsifiable and focused.",
        "Normal",
    )

    # Add revised synthesis after the original conclusion, preserving it.
    anchor = find_paragraph(doc, starts="This draft now contains both the original dense SwiGLU pruning evidence")
    paragraph_after(
        anchor,
        "Version 3 synthesis. The project has progressed from a heuristic spherical score to an exact ellipsoid bound tied to the true RMSNorm feasible set. Controlled experiments show that the new geometry is useful for within-layer selection, while the failed raw ellipsoid allocation reveals that layer sensitivity is a separate problem. At 6.207% expert-width reduction, the ellipsoid selector gives +0.610% WikiText-2 and +1.672% C4 relative PPL under the primary fixed allocation, with paired confidence intervals favoring it over both original-bound and activation rankings. The certificate audit finds zero sampled violations and materially improved tightness. These results justify Version 3 as a serious research manuscript, while the retained historical experiments document the full development path. Final submission should wait for downstream, systems, high-ratio, aggregation, and closest-baseline evidence.",
        "Normal",
    )

    # Add new references immediately before Appendix A.
    target = find_paragraph(doc, exact="Appendix A. Exact Selected MoE Full Benchmark Matrix")
    b = BeforeBuilder(doc, target)
    for ref in [
        "[12] Biao Zhang and Rico Sennrich. Root Mean Square Layer Normalization. NeurIPS, 2019. https://arxiv.org/abs/1910.07467",
        "[13] Ke Li, Zheng Yang, Zhongbin Zhou, Feng Xue, Zhonglin Jiang, and Wenxiao Wang. HEAPr: Hessian-based Efficient Atomic Expert Pruning in Output Space. arXiv:2509.22299, 2025. https://arxiv.org/abs/2509.22299",
        "[14] Yuzhuang Xu, Xu Han, Yuanchi Zhang, Yixuan Wang, Yijun Liu, Shiyu Ji, Qingfu Zhu, and Wanxiang Che. CAMERA: Multi-Matrix Joint Compression for MoE Models via Micro-Expert Redundancy Analysis. arXiv:2508.02322, 2025. https://arxiv.org/abs/2508.02322",
        "[15] Yifu Ding, Jiacheng Wang, Ge Yang, Yongcheng Jing, Jinyang Guo, Xianglong Liu, and Dacheng Tao. Attribution-Guided and Coverage-Maximized Pruning for Structural MoE Compression. arXiv:2606.18304, 2026. https://arxiv.org/abs/2606.18304",
        "[16] Internal benchmark artifact. Qwen3-30B-A3B allocation/ranking target-4 experiment, results/moe_allocation_ranking/target4_n1024_v1, 2026.",
        "[17] Internal benchmark artifact. Qwen3-30B-A3B allocation/ranking target-6 experiment, results/moe_allocation_ranking/target6_rmsnorm_primary_n1024_v1, 2026.",
        "[18] Internal diagnostic artifact. Qwen3-30B-A3B 2×2 allocation/ranking experiment and ellipsoid target-2 validation, 2026.",
        "[19] Kimi Team. Kimi K3: Open Frontier Intelligence. arXiv:2607.24653v2, 2026. https://arxiv.org/abs/2607.24653",
    ]:
        b.p(ref, "Normal")

    # Mathematical appendix, appended after all retained appendices.
    doc.add_heading("Appendix C. Proof of the RMSNorm Ellipsoid Bound", level=1)
    doc.add_paragraph(
        "Let q = x / sqrt(||x||²/d + ε). Then ||q||² = ||x||²/(||x||²/d + ε) < d. With Γ = diag(γ), the post-RMSNorm vector is r = Γq. For one SwiGLU channel, define a = Γg and b = Γu. Its deleted output is c(q) = SiLU(qᵀa)(qᵀb)d_out. Because |SiLU(t)| ≤ |t|,"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("||c(q)|| ≤ |qᵀa| |qᵀb| ||d_out||.").font.name = "Cambria Math"
    doc.add_paragraph(
        "Define M = (abᵀ + baᵀ)/2. Then (qᵀa)(qᵀb) = qᵀMq. The range of M lies in span{a,b}; its two possibly nonzero eigenvalues are λ± = (aᵀb ± ||a||||b||)/2. Hence its spectral radius is ρ(M) = (||a||||b|| + |aᵀb|)/2. The maximum absolute quadratic form over ||q|| ≤ sqrt(d) is dρ(M), attained along an eigenvector corresponding to the eigenvalue of largest magnitude. Multiplying by ||d_out|| proves"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Bᵉˡˡ = d/2 (||Γg|| ||Γu|| + |(Γg)ᵀ(Γu)|) ||d_out||.").font.name = "Cambria Math"
    doc.add_paragraph(
        "The result is exact for the bilinear envelope. Equality for the full SwiGLU channel is not generally expected because |SiLU(t)| can be strictly smaller than |t|, especially on the negative side. This distinction explains why a valid bound can remain loose in empirical routed-token measurements."
    )
    doc.add_heading("Appendix D. Generalization to SiTU-GLU", level=1)
    doc.add_paragraph(
        "For SiTU-GLU, let φ(a) = β₁ tanh(a/β₁)σ(a) and ψ(b) = β₂ tanh(b/β₂). The elementary envelopes |φ(a)| ≤ min(β₁, |a|) and |ψ(b)| ≤ min(β₂, |b|) imply four simultaneous channel bounds: the global cap β₁β₂; a gate cap times a linear up bound β₁ sqrt(d)||Γu||; an up cap times a linear gate bound β₂ sqrt(d)||Γg||; and the ellipsoidal bilinear bound. Taking their minimum and multiplying by the down-column norm gives the composite bound stated in Section 3.2.5. This provides a direct path to bounded-activation MoE architectures while retaining the same pruning unit and physical width reduction."
    )
    doc.add_paragraph(
        "Partial MoE activation and structural pruning address different resources. Top-k routing reduces the number of experts executed per token, but the entire expert pool must still be stored and every activated expert retains its original width. Channel pruning reduces stored parameters and the matrix dimensions of each activated expert. The methods are therefore complementary rather than equivalent."
    )
    doc.add_heading("Appendix E. Version 3 Claim Checklist", level=1)
    checklist = [
        "Supported: the expert-specific ellipsoid formula is a valid upper bound under the stated RMSNorm and SwiGLU assumptions.",
        "Supported on sampled routed inputs: zero numerical violations and improved tightness relative to the spherical bound.",
        "Supported under the target-6 fixed allocation: ellipsoid ranking improves original-bound and activation ranking on WikiText-2 and C4 with paired intervals excluding zero.",
        "Supported directionally at target 4%: ellipsoid ranking improves both tested allocation plans on both PPL corpora.",
        "Not yet supported: a uniform all-expert certificate for p95 aggregation, state-of-the-art high-ratio pruning, downstream-task preservation, or end-to-end serving speedup.",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    # Metadata and update-on-open setting.
    doc.core_properties.title = "Certified Ellipsoidal Bound-Guided Structured SwiGLU and MoE Expert-Channel Pruning — Version 3"
    doc.core_properties.subject = "Working research draft preserving Versions 1–2 and adding the RMSNorm ellipsoid expansion"
    doc.core_properties.comments = "Version 3 internal research manuscript; pending downstream, runtime, max-aggregation, and external-baseline experiments."
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUTPUT)
    print(f"WROTE {OUTPUT}")
    print(f"SHA256 {sha256(OUTPUT)}")


if __name__ == "__main__":
    main()
